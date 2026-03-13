from __future__ import annotations

import io
import json
import os
import shutil
from pathlib import Path
from typing import List

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Chroma

from .config import UPLOAD_CHROMA_DIR, UPLOAD_COLLECTION_NAME
from .ingest import chunk_documents
from .providers import get_embeddings


def _save_uploaded_file(uploaded_file, tmp_dir: Path) -> Path:
    tmp_dir.mkdir(parents=True, exist_ok=True)
    file_path = tmp_dir / uploaded_file.name
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_uploaded_pdf(path: Path) -> List[Document]:
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = path.name
        d.metadata["file_type"] = "pdf"
        d.metadata["source_scope"] = "upload"
    return docs


def _load_uploaded_html(path: Path) -> List[Document]:
    html = _read_text_file(path)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "html",
                "page": 1,
                "source_scope": "upload",
            },
        )
    ]


def _load_uploaded_txt_or_rtf(path: Path) -> List[Document]:
    text = _read_text_file(path)
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": path.suffix.lower().lstrip("."),
                "page": 1,
                "source_scope": "upload",
            },
        )
    ]


def _load_uploaded_csv(path: Path) -> List[Document]:
    df = pd.read_csv(path)
    text = df.to_csv(index=False)
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "csv",
                "page": 1,
                "rows": len(df),
                "columns": ", ".join(map(str, df.columns)),
                "source_scope": "upload",
            },
        )
    ]


def _load_uploaded_json(path: Path) -> List[Document]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    text = json.dumps(data, indent=2, ensure_ascii=False)
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "json",
                "page": 1,
                "source_scope": "upload",
            },
        )
    ]


def _load_uploaded_docx(path: Path) -> List[Document]:
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "docx",
                "page": 1,
                "source_scope": "upload",
            },
        )
    ]


def _load_uploaded_file(path: Path) -> List[Document]:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_uploaded_pdf(path)
    if suffix in {".html", ".htm"}:
        return _load_uploaded_html(path)
    if suffix in {".txt", ".rtf"}:
        return _load_uploaded_txt_or_rtf(path)
    if suffix == ".csv":
        return _load_uploaded_csv(path)
    if suffix == ".json":
        return _load_uploaded_json(path)
    if suffix == ".docx":
        return _load_uploaded_docx(path)

    return []


def build_upload_index(uploaded_files) -> int:
    if not uploaded_files:
        return 0

    embeddings = get_embeddings()
    os.makedirs(UPLOAD_CHROMA_DIR, exist_ok=True)

    tmp_dir = Path("data/tmp_uploads")
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    if Path(UPLOAD_CHROMA_DIR).exists():
        shutil.rmtree(UPLOAD_CHROMA_DIR)
    os.makedirs(UPLOAD_CHROMA_DIR, exist_ok=True)

    docs: List[Document] = []
    for uploaded_file in uploaded_files:
        saved_path = _save_uploaded_file(uploaded_file, tmp_dir)
        docs.extend(_load_uploaded_file(saved_path))

    if not docs:
        return 0

    chunks = chunk_documents(docs)

    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=UPLOAD_CHROMA_DIR,
        collection_name=UPLOAD_COLLECTION_NAME,
    )

    return len(chunks)
