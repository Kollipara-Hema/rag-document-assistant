"""
Load documents from the local repository folder.

This loader is the base document source for the learning lab.
It reads supported files from data/raw_docs and converts them into
LangChain Document objects with metadata.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader

from src.config import RAW_DOCS_DIR


SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".txt", ".csv", ".json", ".docx", ".rtf"}


def _read_text_file(path: Path) -> str:
    """
    Read a text-like file using a forgiving encoding strategy.

    We ignore decoding errors so that imperfect files do not crash
    the whole ingestion step.
    """
    return path.read_text(encoding="utf-8", errors="ignore")


def _load_pdf(path: Path) -> List[Document]:
    """
    Load a PDF file page by page.
    """
    loader = PyPDFLoader(str(path))
    docs = loader.load()

    for doc in docs:
        doc.metadata["source"] = path.name
        doc.metadata["file_type"] = "pdf"
        doc.metadata["source_scope"] = "local"

    return docs


def _load_html(path: Path) -> List[Document]:
    """
    Extract visible text from HTML files.

    We remove script/style content because it is not useful
    for semantic retrieval.
    """
    html = _read_text_file(path)
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    cleaned_text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

    return [
        Document(
            page_content=cleaned_text,
            metadata={
                "source": path.name,
                "file_type": "html",
                "page": 1,
                "source_scope": "local",
            },
        )
    ]


def _load_txt_or_rtf(path: Path) -> List[Document]:
    """
    Load plain text and simple RTF-like files as raw text.

    This is intentionally simple for v1. We can later replace this
    with a richer RTF parser if needed.
    """
    text = _read_text_file(path)

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": path.suffix.lower().lstrip("."),
                "page": 1,
                "source_scope": "local",
            },
        )
    ]


def _load_csv(path: Path) -> List[Document]:
    """
    Convert a CSV into a text representation so it can be chunked
    and indexed like other documents.
    """
    df = pd.read_csv(path)
    csv_text = df.to_csv(index=False)

    return [
        Document(
            page_content=csv_text,
            metadata={
                "source": path.name,
                "file_type": "csv",
                "page": 1,
                "rows": len(df),
                "columns": ", ".join(map(str, df.columns)),
                "source_scope": "local",
            },
        )
    ]


def _load_json(path: Path) -> List[Document]:
    """
    Serialize JSON into pretty-formatted text for retrieval.

    This is not the most advanced JSON indexing approach, but it is
    clean, readable, and useful for educational experiments.
    """
    with open(path, "r", encoding="utf-8") as file:
        data = json.load(file)

    json_text = json.dumps(data, indent=2, ensure_ascii=False)

    return [
        Document(
            page_content=json_text,
            metadata={
                "source": path.name,
                "file_type": "json",
                "page": 1,
                "source_scope": "local",
            },
        )
    ]


def _load_docx(path: Path) -> List[Document]:
    """
    Load text from a DOCX file by concatenating paragraph content.
    """
    doc = DocxDocument(path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "docx",
                "page": 1,
                "source_scope": "local",
            },
        )
    ]


def _load_single_file(path: Path) -> List[Document]:
    """
    Route each file to the correct loader based on file extension.
    """
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix in {".txt", ".rtf"}:
        return _load_txt_or_rtf(path)
    if suffix == ".csv":
        return _load_csv(path)
    if suffix == ".json":
        return _load_json(path)
    if suffix == ".docx":
        return _load_docx(path)

    return []


def load_local_documents(folder_path: Path | None = None) -> List[Document]:
    """
    Load all supported local documents from the repository data folder.

    Parameters
    ----------
    folder_path : Path | None
        Optional override for the local document directory.

    Returns
    -------
    List[Document]
        Loaded documents ready for chunking.
    """
    target_folder = folder_path or RAW_DOCS_DIR

    if not target_folder.exists():
        raise FileNotFoundError(f"Local document folder does not exist: {target_folder}")

    files = sorted(
        [path for path in target_folder.iterdir() if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS]
    )

    if not files:
        raise FileNotFoundError(f"No supported files found in: {target_folder}")

    all_documents: List[Document] = []

    for file_path in files:
        loaded_docs = _load_single_file(file_path)
        all_documents.extend(loaded_docs)

    return all_documents
