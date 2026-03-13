from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config import CHROMA_DIR, COLLECTION_NAME, CHUNK_OVERLAP, CHUNK_SIZE
from .providers import get_embeddings


SUPPORTED_EXTENSIONS = {".pdf", ".html", ".htm", ".txt", ".csv", ".json", ".docx", ".rtf"}


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> List[Document]:
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = path.name
        d.metadata["file_type"] = "pdf"
    return docs


def load_html(path: Path) -> List[Document]:
    html = read_text_file(path)
    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "html",
                "page": 1,
            },
        )
    ]


def load_txt_or_rtf(path: Path) -> List[Document]:
    text = read_text_file(path)
    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": path.suffix.lower().lstrip("."),
                "page": 1,
            },
        )
    ]


def load_csv(path: Path) -> List[Document]:
    df = pd.read_csv(path)
    text = df.to_csv(index=False)

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "csv",
                "page": 1,
                "rows": len(df),
                "columns": list(df.columns),
            },
        )
    ]


def load_json(path: Path) -> List[Document]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    text = json.dumps(data, indent=2, ensure_ascii=False)

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "json",
                "page": 1,
            },
        )
    ]


def load_docx(path: Path) -> List[Document]:
    doc = DocxDocument(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "file_type": "docx",
                "page": 1,
            },
        )
    ]


def load_file(path: Path) -> List[Document]:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf(path)
    if suffix in {".html", ".htm"}:
        return load_html(path)
    if suffix in {".txt", ".rtf"}:
        return load_txt_or_rtf(path)
    if suffix == ".csv":
        return load_csv(path)
    if suffix == ".json":
        return load_json(path)
    if suffix == ".docx":
        return load_docx(path)

    return []


def load_documents(folder: str) -> List[Document]:
    docs: List[Document] = []
    folder_path = Path(folder)

    if not folder_path.exists():
        raise FileNotFoundError(f"Folder not found: {folder}")

    files = sorted(
        [p for p in folder_path.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS]
    )

    if not files:
        raise FileNotFoundError(f"No supported files found in: {folder}")

    for file_path in files:
        loaded = load_file(file_path)
        docs.extend(loaded)

    return docs


def chunk_documents(docs: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    return splitter.split_documents(docs)


def build_chroma_index(chunks: List[Document]) -> None:
    os.makedirs(CHROMA_DIR, exist_ok=True)
    embeddings = get_embeddings()

    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=CHROMA_DIR,
        collection_name=COLLECTION_NAME,
    )


def ingest(folder: str = "data/raw_docs") -> Dict[str, Any]:
    docs = load_documents(folder)
    chunks = chunk_documents(docs)
    build_chroma_index(chunks)

    return {
        "files_loaded": len(set(d.metadata.get("source") for d in docs)),
        "documents_loaded": len(docs),
        "chunks_created": len(chunks),
        "chroma_dir": CHROMA_DIR,
        "collection": COLLECTION_NAME,
    }


if __name__ == "__main__":
    stats = ingest()
    print(stats)
